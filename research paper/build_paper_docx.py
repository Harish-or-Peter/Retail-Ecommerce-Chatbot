"""Convert the markdown draft into a two-column IEEE-style .docx with figures/tables/code.

Run from this folder:
    pip install python-docx
    python build_paper_docx.py
Produces Retail_Ecommerce_Research_Paper.docx next to this script.
"""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
MD   = os.path.join(BASE, "Retail_Ecommerce_Research_Paper_DRAFT.md")
DOCDIR = BASE
OUT  = os.path.join(BASE, "Retail_Ecommerce_Research_Paper.docx")

doc = Document()
normal = doc.styles['Normal']; normal.font.name = 'Times New Roman'; normal.font.size = Pt(10)
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.7)
    sec.left_margin = sec.right_margin = Inches(0.6)

def set_columns(section, num):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:space'), '288')

INLINE = re.compile(r'(\*\*.+?\*\*|\*.+?\*)')
def add_runs(p, text):
    text = re.sub(r'\[([^\]]+)\]\((https?://[^)]+)\)', r'\1 (\2)', text)
    for tok in INLINE.split(text):
        if not tok: continue
        if tok.startswith('**') and tok.endswith('**'):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            p.add_run(tok[1:-1]).italic = True
        else:
            p.add_run(tok)

def add_image(path, caption):
    full = os.path.normpath(os.path.join(DOCDIR, path))
    if os.path.exists(full):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(full, width=Inches(3.2))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
    else:
        doc.add_paragraph(f"[missing image: {path}]")

def add_code(lines):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.left_indent = Inches(0.08); pf.space_before = Pt(3); pf.space_after = Pt(3)
    r = p.add_run("\n".join(lines)); r.font.name = 'Consolas'; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

def add_table(rows):
    ncol = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=ncol); t.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = t.cell(i, j); cell.text = ''
            rr = cell.paragraphs[0].add_run(row[j] if j < len(row) else '')
            rr.font.size = Pt(8)
            if i == 0: rr.bold = True

lines = open(MD, encoding='utf-8').read().split('\n')
i = 0; two_col = False; in_comment = False
while i < len(lines):
    s = lines[i].strip()
    if '<!--' in s:
        in_comment = '-->' not in s; i += 1; continue
    if in_comment:
        if '-->' in s: in_comment = False
        i += 1; continue
    if s == '' or s == '---':
        i += 1; continue
    if s.startswith('## I. Introduction') and not two_col:
        ns = doc.add_section(WD_SECTION.CONTINUOUS)
        ns.top_margin = ns.bottom_margin = Inches(0.7); ns.left_margin = ns.right_margin = Inches(0.6)
        set_columns(ns, 2); two_col = True
    if s.startswith('# ') and not s.startswith('## '):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(s[2:].strip()); r.bold = True; r.font.size = Pt(17); i += 1; continue
    if s.startswith('### '):
        doc.add_heading(s[4:].strip(), level=2); i += 1; continue
    if s.startswith('## '):
        doc.add_heading(s[3:].strip(), level=1); i += 1; continue
    m = re.match(r'!\[(.*?)\]\((.*?)\)', s)
    if m:
        add_image(m.group(2), m.group(1)); i += 1; continue
    if s.startswith('```'):
        block = []; i += 1
        while i < len(lines) and not lines[i].strip().startswith('```'):
            block.append(lines[i]); i += 1
        add_code(block); i += 1; continue
    if s.startswith('|'):
        tb = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            raw = lines[i].strip().strip('|')
            if not re.match(r'^[-:\s|]+$', raw.replace('|', '')):
                tb.append([c.strip() for c in raw.split('|')])
            i += 1
        if tb: add_table(tb)
        continue
    p = doc.add_paragraph(); add_runs(p, s); i += 1

doc.save(OUT)
print("WROTE", OUT)
